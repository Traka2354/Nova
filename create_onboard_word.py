"""
Recreates the "Onboard Experiences – Corfu, Greece" flyer as a Word document.
Layout mirrors the original printed sheet as closely as python-docx allows.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── low-level XML helpers ────────────────────────────────────────────────────

def _shd(cell, fill_hex):
    """Set cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove old shd
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _borders(cell, top=None, bottom=None, left=None, right=None):
    """Apply borders to a cell. Pass a dict or None per side."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in (('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)):
        if val:
            tag = OxmlElement(f'w:{side}')
            tag.set(qn('w:val'),   val.get('val', 'single'))
            tag.set(qn('w:sz'),    val.get('sz',  '12'))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(tag)
        else:
            tag = OxmlElement(f'w:{side}')
            tag.set(qn('w:val'), 'none')
            tcBorders.append(tag)
    tcPr.append(tcBorders)


SOLID = lambda sz='12', color='222222': {'val': 'single', 'sz': sz, 'color': color}
NONE  = None


def box_cell(cell, sz='12', color='222222'):
    s = SOLID(sz, color)
    _borders(cell, top=s, bottom=s, left=s, right=s)


def no_border(cell):
    _borders(cell, top=NONE, bottom=NONE, left=NONE, right=NONE)


def _remove_para_spacing(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


def _set_row_height(row, cm_val):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(qn('w:trHeight')):
        trPr.remove(old)
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(cm_val * 567)))   # 1 cm ≈ 567 twips
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)


def _cell_margin(cell, top=50, bottom=50, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)):
        t = OxmlElement(f'w:{side}')
        t.set(qn('w:w'), str(val))
        t.set(qn('w:type'), 'dxa')
        tcMar.append(t)
    tcPr.append(tcMar)


# ── paragraph / run helpers ──────────────────────────────────────────────────

def add_run(para, text, bold=False, italic=False, size=9,
            color='000000', underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run


def para_in_cell(cell, text='', bold=False, italic=False, size=9,
                 color='000000', underline=False,
                 align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=0, space_after=0,
                 clear_first=False):
    """Add a paragraph to a cell (or reuse first if clear_first=True)."""
    if clear_first:
        p = cell.paragraphs[0]
        for r in p.runs:
            r._element.getparent().remove(r._element)
    else:
        p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        add_run(p, text, bold=bold, italic=italic,
                size=size, color=color, underline=underline)
    return p


# ── section-header helper ────────────────────────────────────────────────────

def section_header_cell(parent_cell, label):
    """
    Insert a 1-row × 1-col nested table whose single cell has a dark
    background and white bold label – replicating the dark header bars.
    """
    nt = parent_cell.add_table(rows=1, cols=1)
    nt.style = 'Table Grid'
    hc = nt.cell(0, 0)
    _shd(hc, '1A1A1A')
    box_cell(hc, sz='8', color='1A1A1A')
    _cell_margin(hc, top=60, bottom=60, left=80, right=80)
    p = hc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _remove_para_spacing(p)
    add_run(p, label, bold=True, size=9, color='FFFFFF')
    return nt


# ── main ─────────────────────────────────────────────────────────────────────

def build_word():
    doc = Document()

    # page setup
    for sec in doc.sections:
        sec.page_width   = Inches(8.5)
        sec.page_height  = Inches(11)
        sec.left_margin  = Cm(1.3)
        sec.right_margin = Cm(1.3)
        sec.top_margin   = Cm(1.2)
        sec.bottom_margin = Cm(1.2)

    # default paragraph spacing
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)

    # ════════════════════════════════════════════════════════════════════
    # TOP SECTION:  title (left) | form fields (right)
    # ════════════════════════════════════════════════════════════════════
    top = doc.add_table(rows=1, cols=2)
    top.style = 'Table Grid'
    top.alignment = WD_TABLE_ALIGNMENT.CENTER
    top.columns[0].width = Inches(3.5)
    top.columns[1].width = Inches(4.0)

    # ── left: title ──────────────────────────────────────────────────
    lc = top.cell(0, 0)
    no_border(lc)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _cell_margin(lc, top=0, bottom=0, left=0, right=80)

    p_title = lc.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _remove_para_spacing(p_title)
    add_run(p_title, 'ONBOARD\nEXPERIENCES', bold=True, size=30)

    p_loc = lc.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_loc.paragraph_format.space_before = Pt(4)
    _remove_para_spacing(p_loc)
    add_run(p_loc, 'CORFU, GREECE', bold=True, size=11)

    p_sub = lc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _remove_para_spacing(p_sub)
    add_run(p_sub, 'APPROACH OUR TEAM FOR ASSISTANCE', bold=True, size=8)

    # ── right: form fields ───────────────────────────────────────────
    rc = top.cell(0, 1)
    no_border(rc)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _cell_margin(rc, top=0, bottom=0, left=80, right=0)

    # clear default paragraph
    rc.paragraphs[0].clear()

    for label in ('GUEST NAME:', 'ROOM NUMBER:',
                  "CREWMEMBER'S NAME & MAPS ID:"):
        ft = rc.add_table(rows=2, cols=1)
        ft.style = 'Table Grid'
        fc = ft.cell(0, 0)
        fc.merge(ft.cell(1, 0))
        box_cell(fc, sz='8')
        _cell_margin(fc, top=60, bottom=60, left=100, right=100)
        fp = fc.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _remove_para_spacing(fp)
        add_run(fp, label, bold=True, size=8)
        # blank line inside box
        fb = fc.add_paragraph()
        _remove_para_spacing(fb)
        # spacer paragraph between boxes
        sp = rc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(4)

    # spacer between top and grid
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(6)
    gap.paragraph_format.space_after  = Pt(0)

    # ════════════════════════════════════════════════════════════════════
    # MAIN GRID: 3 cols
    #   col 0 (35%): Shore Excursion + Airport Transfer
    #   col 1 (2%):  spacer
    #   col 2 (63%): Vibe Beach Club + Thermal Spa  +  [Train | Laundry]
    # ════════════════════════════════════════════════════════════════════
    grid = doc.add_table(rows=3, cols=3)
    grid.style = 'Table Grid'
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER

    # column widths
    W_LEFT   = Inches(2.5)
    W_GAP    = Inches(0.12)
    W_RIGHT  = Inches(4.88)

    for row in grid.rows:
        row.cells[0].width = W_LEFT
        row.cells[1].width = W_GAP
        row.cells[2].width = W_RIGHT

    # hide spacer column borders
    for r in range(3):
        no_border(grid.cell(r, 1))

    # ════════════════════════════════════════════════════════════════════
    # SHORE EXCURSION  (rows 0-1, col 0  merged)
    # ════════════════════════════════════════════════════════════════════
    shore = grid.cell(0, 0).merge(grid.cell(1, 0))
    shore.width = W_LEFT
    box_cell(shore)
    _shd(shore, 'FFFFFF')
    shore.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _cell_margin(shore, top=0, bottom=80, left=0, right=0)

    section_header_cell(shore, 'SHORE EXCURSION')

    def shore_block(title, items):
        pt = shore.add_paragraph()
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt.paragraph_format.space_before = Pt(5)
        _remove_para_spacing(pt)
        add_run(pt, title, underline=True, size=9)
        for line, is_price in items:
            pl = shore.add_paragraph()
            pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _remove_para_spacing(pl)
            add_run(pl, line, bold=is_price, size=8)

    shore_block('Kotor Tours', [
        ('BOKA BAY CATAMARAN SAIL &', False),
        ('BEACH', False),
        ('$199.99Adult | $149.99 Child', True),
        ('HISTORIC KOTOR & BUDVA', False),
        ('$99.99Adult | $79.99 Child', True),
    ])
    shore_block('Split Tours', [
        ('TROGIR & BEACH', False),
        ('$89.99Adult | $49.99 Child', True),
        ('SPLIT OPEN-AIR SIGHTSEEING', False),
        ('& WALKING TOUR', False),
        ('$69.99Adult | $49.99 Child', True),
    ])
    pv = shore.add_paragraph()
    pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pv.paragraph_format.space_before = Pt(6)
    _remove_para_spacing(pv)
    add_run(pv, 'Visit Shore Excursions Desk', italic=True, size=8)
    pv2 = shore.add_paragraph()
    pv2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _remove_para_spacing(pv2)
    add_run(pv2, 'Deck 7, Mid', italic=True, size=8)

    # ════════════════════════════════════════════════════════════════════
    # VIBE BEACH CLUB  (row 0, col 2)
    # ════════════════════════════════════════════════════════════════════
    vibe = grid.cell(0, 2)
    vibe.width = W_RIGHT
    box_cell(vibe)
    _shd(vibe, 'FFFFFF')
    vibe.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _cell_margin(vibe, top=0, bottom=80, left=0, right=0)

    section_header_cell(vibe, 'VIBE BEACH CLUB')

    for text, bold, size in [
        ('Escape to luxury at sea.', False, 8),
        ('Unwind in style with an exclusive, adults-only retreat', False, 8),
        ('with ocean views & enjoy premium touches designed', False, 8),
        ('just for you.', False, 8),
        ('Limited access available.', False, 8),
        ('Reserve early for the ultimate experience.', False, 8),
    ]:
        pv = vibe.add_paragraph()
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pv.paragraph_format.space_before = Pt(2)
        _remove_para_spacing(pv)
        add_run(pv, text, bold=bold, size=size)

    for text, bold, size in [
        ('Vibe Cruise Pass – $209/person (4 days)', True, 9),
        ('Vibe Cabanas – $499/2pax (4 days)', True, 9),
    ]:
        pv = vibe.add_paragraph()
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pv.paragraph_format.space_before = Pt(3)
        _remove_para_spacing(pv)
        add_run(pv, text, bold=bold, size=size)

    pv = vibe.add_paragraph()
    pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pv.paragraph_format.space_before = Pt(4)
    _remove_para_spacing(pv)
    add_run(pv, 'Visit Guest Service Desk, Deck 7 Midship', italic=True, size=8)

    # ════════════════════════════════════════════════════════════════════
    # THERMAL SPA EXPERIENCE  (row 1, col 2)
    # ════════════════════════════════════════════════════════════════════
    spa = grid.cell(1, 2)
    spa.width = W_RIGHT
    box_cell(spa)
    _shd(spa, 'FFFFFF')
    spa.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _cell_margin(spa, top=0, bottom=80, left=0, right=0)

    section_header_cell(spa, 'THERMAL SPA EXPERIENCE')

    spa_desc = ('Discover a dedicated sanctuary for relaxation & rejuvenation, '
                'thoughtfully redesigned with enhanced amenities to elevate '
                'your wellness journey at sea. Unwind all-day with exclusive '
                'adults-only access to Mandara Spa\'s serene Thermal Suite '
                'experience.')
    ps = spa.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps.paragraph_format.space_before = Pt(4)
    _remove_para_spacing(ps)
    add_run(ps, spa_desc, size=8)

    ps2 = spa.add_paragraph()
    ps2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps2.paragraph_format.space_before = Pt(6)
    _remove_para_spacing(ps2)
    add_run(ps2, 'for only $89/person (day pass)', bold=True, size=13)

    ps3 = spa.add_paragraph()
    ps3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ps3.paragraph_format.space_before = Pt(3)
    _remove_para_spacing(ps3)
    add_run(ps3, 'from 8am to 10pm | Deck 12 Forward, Spa', italic=True, size=8)

    # ════════════════════════════════════════════════════════════════════
    # AIRPORT TRANSFER  (row 2, col 0)
    # ════════════════════════════════════════════════════════════════════
    air = grid.cell(2, 0)
    air.width = W_LEFT
    box_cell(air)
    _shd(air, 'FFFFFF')
    air.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _cell_margin(air, top=0, bottom=80, left=0, right=0)

    section_header_cell(air, 'AIRPORT TRANSFER')

    for text, bold, italic, size in [
        ('From Porto di Ravenna to the', False, False, 8),
        ('following airports:', False, False, 8),
        ('', False, False, 5),
        ('-Venice Marco Polo Airport', False, False, 8),
        ('(VCE)', False, False, 8),
        ('-Bologna Airport (BLQ)', False, False, 8),
        ('-Forli International Airport', False, False, 8),
        ('(FRL)', False, False, 8),
        ('', False, False, 5),
        ('for only $99.50/person', True, False, 11),
        ('until 10am of May 23 only', False, True, 8),
        ('', False, False, 5),
        ('Visit Guest Services Desk,', False, True, 8),
        ('Deck 7 Mid', False, True, 8),
    ]:
        pa = air.add_paragraph()
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _remove_para_spacing(pa)
        add_run(pa, text, bold=bold, italic=italic, size=size)

    # ════════════════════════════════════════════════════════════════════
    # BOTTOM-RIGHT: Train Station Transfer | Laundry Promo
    # ════════════════════════════════════════════════════════════════════
    bot_right = grid.cell(2, 2)
    bot_right.width = W_RIGHT
    no_border(bot_right)
    _shd(bot_right, 'FFFFFF')
    # clear default paragraph
    bot_right.paragraphs[0].clear()

    # nested 1×2 table inside bot_right
    inner = bot_right.add_table(rows=1, cols=2)
    inner.style = 'Table Grid'
    inner.columns[0].width = Inches(2.4)
    inner.columns[1].width = Inches(2.48)

    # ── Train Station Transfer ────────────────────────────────────────
    tr_cell = inner.cell(0, 0)
    tr_cell.width = Inches(2.4)
    box_cell(tr_cell)
    _shd(tr_cell, 'FFFFFF')
    tr_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _cell_margin(tr_cell, top=0, bottom=80, left=0, right=0)

    section_header_cell(tr_cell, 'TRAIN STATION TRANSFER')

    for text, bold, italic, size in [
        ('From Porto di Ravenna to', False, False, 8),
        ('Ravenna Train Station', False, False, 8),
        ('', False, False, 5),
        ('for only $23.00/person', True, False, 12),
        ('', False, False, 5),
        ('available until 12nn of May 22', False, False, 8),
        ('', False, False, 5),
        ('Visit Guest Services Desk,', False, True, 8),
        ('Deck 7 Mid', False, True, 8),
    ]:
        pt = tr_cell.add_paragraph()
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _remove_para_spacing(pt)
        add_run(pt, text, bold=bold, italic=italic, size=size)

    # ── Laundry Promo ────────────────────────────────────────────────
    lau_cell = inner.cell(0, 1)
    lau_cell.width = Inches(2.48)
    box_cell(lau_cell)
    _shd(lau_cell, 'FFFFFF')
    lau_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _cell_margin(lau_cell, top=0, bottom=80, left=0, right=0)

    section_header_cell(lau_cell, 'LAUNDRY PROMO')

    for text, bold, italic, size in [
        ('', False, False, 5),
        ('Leave the laundry,', True, False, 10),
        ('Seize the day!', True, False, 10),
        ('Fill the bag–we\'ll handle', False, False, 8),
        ('the rest! Wash & fold in', False, False, 8),
        ('48 hours', False, False, 8),
        ('', False, False, 5),
        ('for only $39.50', True, False, 12),
    ]:
        pl = lau_cell.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _remove_para_spacing(pl)
        add_run(pl, text, bold=bold, italic=italic, size=size)

    # ════════════════════════════════════════════════════════════════════
    out = '/home/user/sonic-ai-copyx/Onboard_Experiences.docx'
    doc.save(out)
    print(f'Saved: {out}')


build_word()
